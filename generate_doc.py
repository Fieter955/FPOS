import os
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: Library 'python-docx' tidak ditemukan.")
    print("Silakan install dengan menjalankan: pip install python-docx")
    exit(1)

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def generate_documentation():
    doc = Document()

    # --- HALAMAN JUDUL ---
    title = doc.add_heading('\n\n\n\nDokumentasi Teknis & SDLC\nSistem FPOS (Point of Sale & ERP)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Versi 1.0 - Juni 2026\nProfesional & Terintegrasi')
    run.font.size = Pt(14)
    
    doc.add_page_break()

    # --- PENDAHULUAN ---
    add_heading(doc, '1. Pendahuluan', 1)
    doc.add_paragraph(
        "FPOS adalah sistem manajemen Point of Sale (POS) dan Enterprise Resource Planning (ERP) "
        "sederhana namun kuat yang dirancang untuk menangani operasional retail, stok barang, "
        "hingga laporan akuntansi otomatis. Sistem ini mendukung multi-cabang dan integrasi "
        "real-time antara transaksi kasir dengan jurnal keuangan."
    )

    # --- TEKNOLOGI STACK ---
    add_heading(doc, '2. Arsitektur & Teknologi', 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Komponen'
    hdr_cells[1].text = 'Teknologi'
    
    tech_data = [
        ('Frontend', 'Vanilla JavaScript, HTML5, CSS Variables'),
        ('Backend', 'FastAPI (Python 3.x)'),
        ('Database', 'SQLite (SQLAlchemy ORM)'),
        ('Auth', 'JWT (JSON Web Token) & Bcrypt'),
        ('Reporting', 'Excel Export (openpyxl), PDF Receipt'),
    ]
    
    for comp, tech in tech_data:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = tech

    # --- SDLC ---
    add_heading(doc, '3. Software Development Life Cycle (SDLC)', 1)
    
    sdlc_phases = [
        ('Analisis Kebutuhan', 
         'Identifikasi kebutuhan sistem POS yang terintegrasi dengan akuntansi (ERP). '
         'Mendukung fitur multi-cabang, manajemen stok, hutang-piutang, dan pelaporan keuangan.'),
        
        ('Desain Sistem', 
         'Perancangan skema database menggunakan SQLAlchemy. Pembuatan UI modular dengan '
         'reusable components di JavaScript. Pendefinisian Pydantic schemas untuk validasi data API.'),
        
        ('Implementasi', 
         'Pengembangan API menggunakan FastAPI dengan routing modular. Pembuatan frontend '
         'yang responsif tanpa framework berat untuk memastikan performa tinggi di perangkat low-end.'),
        
        ('Pengujian (Testing)', 
         'Validasi alur transaksi dari POS hingga masuk ke Jurnal Umum secara otomatis. '
         'Uji coba sinkronisasi stok antar cabang dan validasi laporan Laba Rugi serta Neraca.'),
        
        ('Deployment', 
         'Pengemasan aplikasi menggunakan PyInstaller untuk menjadi executable tunggal agar '
         'memudahkan instalasi di sisi klien tanpa perlu setup Python manual.'),
        
        ('Pemeliharaan', 
         'Pemantauan sistem melalui Audit Logs. Backup database ipos.db secara berkala. '
         'Pembaruan fitur berdasarkan feedback operasional di lapangan.')
    ]

    for phase, desc in sdlc_phases:
        add_heading(doc, phase, 2)
        doc.add_paragraph(desc)

    # --- API ENDPOINTS ---
    add_heading(doc, '4. Dokumentasi API Endpoints', 1)
    doc.add_paragraph("Daftar endpoint API utama yang tersedia di sistem backend FPOS:")

    def add_api_table(section_title, endpoints):
        add_heading(doc, section_title, 2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Method'
        hdr_cells[1].text = 'Endpoint'
        hdr_cells[2].text = 'Deskripsi'
        
        for method, path, desc in endpoints:
            row_cells = table.add_row().cells
            row_cells[0].text = method
            row_cells[1].text = path
            row_cells[2].text = desc

    # Auth
    add_api_table('Authentication', [
        ('POST', '/api/auth/login', 'Autentikasi user dan pemberian token JWT'),
        ('POST', '/api/auth/register', 'Pendaftaran user baru (Admin Only)'),
        ('GET', '/api/auth/me', 'Mendapatkan data user yang sedang login'),
        ('GET', '/api/auth/users', 'List semua user sistem'),
        ('GET', '/api/auth/audit-log', 'Melihat log aktivitas sistem'),
    ])

    # Items
    add_api_table('Master Data (Items)', [
        ('GET', '/api/items/', 'Mendapatkan daftar barang (mendukung filter)'),
        ('POST', '/api/items/', 'Menambah barang baru'),
        ('PUT', '/api/items/{id}', 'Update data barang'),
        ('DELETE', '/api/items/{id}', 'Hapus barang'),
        ('POST', '/api/items/import/excel', 'Bulk import barang dari file Excel'),
    ])

    # Sales
    add_api_table('Sales & POS', [
        ('GET', '/api/sales/', 'List riwayat penjualan'),
        ('POST', '/api/sales/', 'Mencatat transaksi penjualan baru'),
        ('GET', '/api/sales/print/{id}', 'Mendapatkan data untuk struk penjualan'),
        ('POST', '/api/sales/cancel/{id}', 'Membatalkan transaksi penjualan'),
    ])

    # Accounting
    add_api_table('Accounting & Finance', [
        ('GET', '/api/accounting/accounts/', 'Daftar Chart of Accounts (CoA)'),
        ('GET', '/api/accounting/journals/', 'List Jurnal Umum'),
        ('GET', '/api/accounting/ledger/', 'Buku Besar per akun'),
        ('GET', '/api/accounting/income-statement', 'Laporan Laba Rugi'),
        ('GET', '/api/accounting/balance-sheet', 'Laporan Neraca (Posisi Keuangan)'),
    ])

    # --- STANDAR UI/UX ---
    add_heading(doc, '5. Standar Komponen UI', 1)
    doc.add_paragraph(
        "Sistem menggunakan komponen UI yang konsisten untuk memastikan pengalaman pengguna yang seragam:"
    )
    ui_rules = [
        "Searchable Dropdowns: Menggunakan createPremiumCombo untuk semua input pilihan.",
        "Unified Grid: Tabel pembelian dan penjualan yang konsisten (createPurchaseGrid).",
        "Payment Modal: Modul pembayaran terpadu untuk Cash/Bank/Split Payment.",
        "Barcode Scanner: Hook global untuk mendeteksi input scanner hardware otomatis."
    ]
    for rule in ui_rules:
        doc.add_paragraph(rule, style='List Bullet')

    # Save
    output_path = "Dokumentasi_Sistem_FPOS.docx"
    doc.save(output_path)
    print(f"Sukses! Dokumentasi telah dibuat: {output_path}")

if __name__ == "__main__":
    generate_documentation()
